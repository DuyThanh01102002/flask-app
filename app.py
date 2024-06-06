from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import io
import openpyxl
import tempfile

app = Flask(__name__)

def replace_text_in_paragraph(paragraph, data):
    for key, value in data.items():
        if key in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if key in inline[i].text:
                    if key == '[wr]':
                        value += '\n'
                    inline[i].text = inline[i].text.replace(key, value.strip())
                    if key == '[dob]':
                        inline[i].italic = True
                    if key in ['[from]', '[to]']:
                        inline[i].bold = True

def format_dob(dob_str):
    try:
        if not dob_str:
            return ""
        parts = dob_str.split('/')
        if len(parts) == 3:
            day, month, year = parts
            day = int(day) if day else 1
            month = int(month) if month else 1
            year = int(year)
            dob_date = datetime(year, month, day)
            return f"{dob_date.day} tháng {dob_date.month} năm {dob_date.year}"
        elif len(parts) == 2:
            month, year = parts
            month = int(month) if month else 1
            year = int(year)
            dob_date = datetime(year, month, 1)
            return f"    tháng {dob_date.month} năm {dob_date.year}"
        elif len(parts) == 1:
            year = int(parts[0])
            return f"    tháng   năm {year}"
        else:
            return dob_str
    except ValueError:
        return dob_str

def format_gift(gift_text):
    if not gift_text:
        return ""
    lines = gift_text.split('\n')
    if len(lines) == 1:
        return f"- {lines[0].strip()}"
    formatted_lines = ["- Như trên;"] + [f"- {line.strip()}" for line in lines if line.strip()]
    return '\n'.join(formatted_lines)

def add_bulleted_paragraph(document, text):
    p = document.add_paragraph(text)
    p.style = 'List Bullet'
    return p

def fill_template(template_path, output_path, data):
    if os.path.exists(output_path):
        doc = Document(output_path)
    else:
        doc = Document(template_path)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)

    if '[dob]' in data:
        data['[dob]'] = format_dob(data['[dob]'])
    if '[gift]' in data:
        data['[gift]'] = format_gift(data['[gift]'])

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)

    doc.save(output_path)

def create_excel_file(output_path, data):
    # Kiểm tra xem tệp đã tồn tại chưa
    if os.path.exists(output_path):
        wb = openpyxl.load_workbook(output_path)
        ws = wb.active
        start_row = ws.max_row + 1  # Bắt đầu từ hàng tiếp theo
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        # Tạo tiêu đề cột
        headers = ['Số thứ tự', 'Về việc', 'Ngày', 'Kính gửi', 'Nội dung', 'Nơi nhận', 'Tên vị trí', 'Tên']
        ws.append(headers)
        start_row = 2  # Dòng bắt đầu dữ liệu (sau tiêu đề)

    # Tạo một hàng mới với dữ liệu và số thứ tự tự động
    new_row = [start_row - 1]  # Số thứ tự bắt đầu từ 1
    new_row += [data.get('Về việc'), data.get('Ngày'), data.get('Kính gửi'), data.get('Nội dung'), data.get('Nơi nhận'), data.get('Tên vị trí'), data.get('Tên')]
    
    # Thêm hàng mới vào trang tính
    ws.append(new_row)

    # Lưu tệp Excel
    wb.save(output_path)
@app.route('/form1', methods=['GET', 'POST'])
def form1():
    if request.method == 'POST':
        data = {
            '[wr]': request.form.get('wr'),
            '[dob]': request.form.get('dob'),
            '[send]': request.form.get('send'),
            '[cont]': request.form.get('cont'),
            '[gift]': request.form.get('gift'),
            '[from]': request.form.get('from'),
            '[to]': request.form.get('to')
        }
        cont_text = request.form.get('cont').replace('\r', '\n')
        cont_lines = cont_text.split('\n')
        formatted_cont = cont_lines[0]
        if len(cont_lines) > 1:
            for line in cont_lines[1:]:
                formatted_cont += line.strip() + '\t\n\t'
        else:
            formatted_cont += '\t'
        formatted_cont = formatted_cont.replace('\t\n\t\t\n\t', '\t\n\t')
        data['[cont]'] = formatted_cont
        print(f"Received data: {data}")

        template_path = r'D:\form\test.docx'
        output_docx_path = r'D:\formkhucnc.docx'

        if not os.path.exists(template_path):
            return f"Template file not found at {template_path}"

        try:
            fill_template(template_path, output_docx_path, data)
        except Exception as e:
            print(f"An error occurred: {e}")
            return f"An error occurred: {e}"

        return render_template('formmau.html', form_data=request.form)
    
    return render_template('formmau.html')

@app.route('/form2', methods=['GET', 'POST'])
def form2():
    if request.method == 'POST':
        data = {
            '[wr]': request.form.get('wr'),
            '[dob]': request.form.get('dob'),
            '[send]': request.form.get('send'),
            '[cont]': request.form.get('cont'),
            '[gift]': request.form.get('gift'),
            '[from]': request.form.get('from'),
            '[to]': request.form.get('to')
        }
        cont_text = request.form.get('cont').replace('\r', '\n')
        cont_lines = cont_text.split('\n')
        formatted_cont = cont_lines[0]
        if len(cont_lines) > 1:
            for line in cont_lines[1:]:
                formatted_cont += line.strip() + '\t\n\t'
        else:
            formatted_cont += '\t'
        formatted_cont = formatted_cont.replace('\t\n\t\t\n\t', '\t\n\t')
        data['[cont]'] = formatted_cont
        print(f"Received data: {data}")

        template_path = r'D:\form\maupc.docx'
        output_docx_path = r'D:\formphieuchuyen.docx'

        if not os.path.exists(template_path):
            return f"Template file not found at {template_path}"

        try:
            fill_template(template_path, output_docx_path, data)
        except Exception as e:
            print(f"An error occurred: {e}")
            return f"An error occurred: {e}"

        return render_template('formpchuyen.html', form_data=request.form)
    
    return render_template('formpchuyen.html')

@app.route('/export_excel', methods=['POST'])
def export_excel():
    data = {
        'Về việc': request.form.get('wr'),
        'Ngày': request.form.get('dob'),
        'Kính gửi': request.form.get('send'),
        'Nội dung': request.form.get('cont'),
        'Nơi nhận': request.form.get('gift'),
        'Tên vị trí': request.form.get('from'),
        'Tên': request.form.get('to')
    }
    print(f"Received data for Excel export: {data}")

    output_xlsx_path = r'D:\formphieuchuyen.xlsx'

    try:
        create_excel_file(output_xlsx_path, data)  # Truyền đường dẫn trực tiếp vào hàm create_excel_file()
        print(f"Excel file saved at {output_xlsx_path}")
        return send_file(output_xlsx_path, as_attachment=True)
    except Exception as e:
        print(f"An error occurred while exporting Excel for Form 1: {e}")
        return f"An error occurred while exporting Excel for Form 1: {e}"

@app.route('/export_excel_form1', methods=['POST'])
def export_excel_form1():
    data = {
        'Về việc': request.form.get('wr'),
        'Ngày': request.form.get('dob'),
        'Kính gửi': request.form.get('send'),
        'Nội dung': request.form.get('cont'),
        'Nơi nhận': request.form.get('gift'),
        'Tên vị trí': request.form.get('from'),
        'Tên': request.form.get('to')
    }
    print(f"Received data for Form 1 Excel export: {data}")

    output_xlsx_path = r'D:\formkhucnc.xlsx'

    try:
        
        create_excel_file(output_xlsx_path, data)  # Truyền đường dẫn trực tiếp vào hàm create_excel_file()
        print(f"Excel file saved at {output_xlsx_path}")
        return send_file(output_xlsx_path, as_attachment=True)
    except Exception as e:
        print(f"An error occurred while exporting Excel for Form 1: {e}")
        return f"An error occurred while exporting Excel for Form 1: {e}"
@app.route('/')
def index():
    return render_template('chonform.html')
if __name__ == '__main__':
    app.run(debug=True)